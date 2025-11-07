# -*- coding: utf-8 -*-
from odoo import models, fields, api
from datetime import date
import base64
import io
import re
import logging

_logger = logging.getLogger(__name__)


class LegalRegulation(models.Model):
    _name = 'legal.regulation'
    _description = 'Legal Regulation'
    _inherit = ['mail.thread', 'mail.activity.mixin']
    _order = 'hierarchy_order, tahun desc, nomor'
    _rec_name = 'judul'
    
    # Informasi Dasar - Diurutkan berdasarkan hierarki hukum Indonesia
    tipe_dokumen = fields.Selection([
        ('uud_1945', 'Undang-undang Dasar 1945'),
        ('tap_mpr', 'Ketetapan MPR'),
        ('undang_undang', 'Undang-Undang'),
        ('perpu', 'Peraturan Pemerintah Pengganti Undang-Undang'),
        ('peraturan_pemerintah', 'Peraturan Pemerintah'),
        ('peraturan_presiden', 'Peraturan Presiden'),
        ('keputusan_presiden', 'Keputusan Presiden'),
        ('instruksi_presiden', 'Instruksi Presiden'),
        ('peraturan_menteri', 'Peraturan Menteri'),
        ('keputusan_menteri', 'Keputusan Menteri'),
        ('peraturan_daerah', 'Peraturan Daerah'),
        ('peraturan_gubernur', 'Peraturan Gubernur'),
    ], string='Tipe Dokumen', required=True, default='undang_undang')
    
    # Field untuk menentukan urutan hierarki
    hierarchy_order = fields.Integer('Hierarchy Order', compute='_compute_hierarchy_order', store=True)
    
    judul = fields.Text('Judul', required=True, help='Judul lengkap peraturan')
    
    teu = fields.Char('T.E.U. (Tempat Terbit/Entitas Unit)', 
                      help='Contoh: Indonesia, Kementerian Pemberdayaan Perempuan dan Perlindungan Anak',
                      default='Indonesia')
    
    nomor = fields.Char('Nomor', required=True, help='Nomor peraturan')
    
    # Bentuk Peraturan
    bentuk = fields.Char('Bentuk', required=True, 
                         help='Contoh: Peraturan Menteri Pemberdayaan Perempuan Dan Perlindungan Anak')
    
    bentuk_singkat = fields.Char('Bentuk Singkat', 
                                 help='Contoh: Permen PPPA')
    
    # Tanggal dan Waktu
    tahun = fields.Integer('Tahun', required=True, default=lambda self: date.today().year)
    
    tempat_penetapan = fields.Char('Tempat Penetapan', default='Jakarta')
    
    tanggal_penetapan = fields.Date('Tanggal Penetapan', required=True)
    
    tanggal_pengundangan = fields.Date('Tanggal Pengundangan')
    
    tanggal_berlaku = fields.Date('Tanggal Berlaku')
    
    # Informasi Sumber dan Status
    sumber = fields.Text('Sumber', 
                         help='Contoh: BN 2025 (314); 70 hlm')
    
    subjek = fields.Text('Subjek', 
                         help='Contoh: KELUARGA, PERLINDUNGAN ANAK, PEREMPUAN / WANITA')
    
    status = fields.Selection([
        ('berlaku', 'Berlaku'),
        ('dicabut', 'Dicabut'),
        ('diubah', 'Diubah'),
        ('ditunda', 'Ditunda'),
        ('tidak_berlaku', 'Tidak Berlaku'),
    ], string='Status', required=True, default='berlaku')
    
    bahasa = fields.Selection([
        ('bahasa_indonesia', 'Bahasa Indonesia'),
        ('bahasa_inggris', 'Bahasa Inggris'),
        ('bahasa_daerah', 'Bahasa Daerah'),
    ], string='Bahasa', default='bahasa_indonesia')
    
    lokasi = fields.Char('Lokasi', 
                         help='Contoh: Kementerian Pemberdayaan Perempuan dan Perlindungan Anak')
    
    bidang = fields.Selection([
        ('hukum_pidana', 'HUKUM PIDANA'),
        ('hukum_perdata', 'HUKUM PERDATA'),
        ('hukum_administrasi_negara', 'HUKUM ADMINISTRASI NEGARA'),
        ('hukum_tata_negara', 'HUKUM TATA NEGARA'),
        ('hukum_internasional', 'HUKUM INTERNASIONAL'),
        ('hukum_bisnis', 'HUKUM BISNIS'),
        ('hukum_keluarga', 'HUKUM KELUARGA'),
        ('hukum_lingkungan', 'HUKUM LINGKUNGAN'),
        ('hukum_tenaga_kerja', 'HUKUM TENAGA KERJA'),
        ('hukum_pajak', 'HUKUM PAJAK'),
    ], string='Bidang Hukum')
    
    # Field Tambahan
    active = fields.Boolean('Active', default=True)
    
    keterangan = fields.Text('Keterangan Tambahan')
    
    # Field untuk Upload File PDF/TXT
    file_pdf = fields.Binary('File PDF', 
                             attachment=True,
                             help='File PDF peraturan (generated otomatis dari DOCX atau upload manual)')
    file_docx = fields.Binary('File DOCX (Word)', 
                              attachment=True,
                              help='Upload file DOCX peraturan (lebih akurat untuk extract text)')
    file_txt = fields.Binary('File TXT (Plain Text)',
                             attachment=True,
                             help='Upload file TXT (plain text) untuk ekstraksi teks yang lebih sederhana dan stabil')
    file_name = fields.Char('Nama File')
    file_size = fields.Integer('Ukuran File (KB)', compute='_compute_file_size', store=True)
    
    # Field untuk Full-Text Search
    isi_peraturan = fields.Html('Isi Peraturan', 
                               default='<p>Isi peraturan belum tersedia</p>',
                               help='Konten lengkap peraturan untuk pencarian mendalam')
    kata_kunci = fields.Text('Kata Kunci', 
                            default='',
                            help='Kata kunci tambahan untuk pencarian')
    ringkasan = fields.Text('Ringkasan', 
                           default='Ringkasan belum tersedia',
                           help='Ringkasan singkat peraturan')
    
    # Computed Fields
    nama_lengkap = fields.Char('Nama Lengkap', compute='_compute_nama_lengkap', store=True)
    
    is_berlaku_aktif = fields.Boolean('Berlaku Aktif', compute='_compute_is_berlaku_aktif')
    
    @api.depends('bentuk_singkat', 'nomor', 'tahun')
    def _compute_nama_lengkap(self):
        for record in self:
            if record.bentuk_singkat and record.nomor and record.tahun:
                record.nama_lengkap = f"{record.bentuk_singkat} No. {record.nomor} Tahun {record.tahun}"
            else:
                record.nama_lengkap = record.judul or 'Peraturan Baru'
    
    @api.depends('status', 'tanggal_berlaku')
    def _compute_is_berlaku_aktif(self):
        today = date.today()
        for record in self:
            if record.status == 'berlaku':
                if record.tanggal_berlaku:
                    record.is_berlaku_aktif = record.tanggal_berlaku <= today
                else:
                    record.is_berlaku_aktif = True
            else:
                record.is_berlaku_aktif = False
    
    @api.depends('file_pdf', 'file_docx', 'file_txt')
    def _compute_file_size(self):
        """Hitung ukuran file dalam KB"""
        for record in self:
            if record.file_txt:
                record.file_size = int(len(record.file_txt) / 1.37 / 1024)
            elif record.file_docx:
                # Priority DOCX
                record.file_size = int(len(record.file_docx) / 1.37 / 1024)
            elif record.file_pdf:
                record.file_size = int(len(record.file_pdf) / 1.37 / 1024)
            else:
                record.file_size = 0
    
    def _fix_broken_words(self, text):
        """Fix words that are broken with spaces in the middle"""
        import re
        
        # Common Indonesian word patterns that might be broken
        
        # Pattern 1: Fix broken words with 1-3 letter chunks separated by single space
        def merge_short_chunks(text):
            # Find patterns like "xxx xx" or "xx xxx" where total length > 4
            pattern = r'\b([a-zA-Z]{2,4})\s([a-zA-Z]{2,5})\b'
            
            def check_and_merge(match):
                word1 = match.group(1)
                word2 = match.group(2)
                combined = word1 + word2
                
                # Common Indonesian prefixes that indicate word should be merged
                common_prefixes = ['per', 'ber', 'ter', 'men', 'mem', 'pen', 'pem', 'di', 'ke', 'se']
                # Common Indonesian suffixes
                common_suffixes = ['an', 'kan', 'nya', 'ku', 'mu', 'lah', 'kah', 'tah']
                
                # If word1 is a common prefix and combined makes sense
                if word1.lower() in common_prefixes:
                    return combined
                
                # If word2 is a common suffix
                if word2.lower() in common_suffixes:
                    return combined
                
                # If both are very short (2-3 chars) and combined length > 5, likely broken
                if len(word1) <= 3 and len(word2) <= 3 and len(combined) >= 5:
                    return combined
                
                # Otherwise keep original with space
                return match.group(0)
            
            return re.sub(pattern, check_and_merge, text)
        
        # Pattern 2: Fix broken words in ALL CAPS (existing fix)
        def fix_caps_words(match):
            word = match.group(0)
            if len(word.replace(' ', '')) > 5:
                return word.replace(' ', '')
            return word
        
        text = re.sub(r'\b([A-Z]\s){3,}[A-Z]\b', fix_caps_words, text)
        
        # Pattern 3: Fix specific common broken patterns
        # "xxx xx xxx" patterns (3+ segments)
        pattern_multi = r'\b([a-zA-Z]{2,4})\s([a-zA-Z]{1,3})\s([a-zA-Z]{2,5})\b'
        
        def merge_three_chunks(match):
            combined = match.group(1) + match.group(2) + match.group(3)
            # If middle segment is very short (1-2 chars), likely broken
            if len(match.group(2)) <= 2 and len(combined) >= 6:
                return combined
            return match.group(0)
        
        text = re.sub(pattern_multi, merge_three_chunks, text)
        
        # Apply the merge_short_chunks multiple times to catch nested patterns
        for _ in range(2):  # Run twice to catch consecutive broken words
            text = merge_short_chunks(text)
        
        return text
    
    def _format_text_to_html(self, text):
        """Format extracted text into properly structured HTML"""
        import re
        
        if not text or not text.strip():
            return ""
        
        # Fix broken words first
        text = self._fix_broken_words(text)
        
        lines = text.split('\n')
        html_parts = []
        current_paragraph = []
        in_list = False
        
        # Patterns untuk mendeteksi struktur
        # Pattern untuk numbering
        numbering_pattern = re.compile(r'^\s*(\d+\.|[a-z]\.|[A-Z]\.|[ivxIVX]+\.|[(\d]+[\).]|[(\w]+[\).])\s+(.+)$')
        # Pattern untuk BAB, Pasal, Ayat, dll
        section_pattern = re.compile(r'^\s*(BAB|PASAL|Pasal|Ayat|Bagian|Paragraf|BAGIAN|PARAGRAF)\s+(.+)$', re.IGNORECASE)
        # Pattern untuk header (ALL CAPS line)
        header_pattern = re.compile(r'^[A-Z\s]{10,}$')
        # Pattern untuk Mengingat:, Menimbang:, dll
        legal_intro_pattern = re.compile(r'^\s*(Mengingat|Menimbang|Memperhatikan|Menetapkan|Memutuskan|Dengan Rahmat)\s*:\s*(.*)$', re.IGNORECASE)
        
        # State for penjelasan block rendering
        in_penjelasan_block = False
        penjelasan_lines = []
        penjelasan_header_pat = re.compile(r'^\s*💡 \[Penjelasan\s+(Pasal|Ayat|Huruf)[^\]]*\]:\s*$', re.IGNORECASE)

        def flush_current_paragraph():
            nonlocal current_paragraph, in_list
            if current_paragraph:
                if isinstance(current_paragraph[0], tuple):
                    para_text = ' '.join([item[0] if isinstance(item, tuple) else item for item in current_paragraph])
                    indent_px = current_paragraph[0][1] if isinstance(current_paragraph[0], tuple) else 0
                else:
                    para_text = ' '.join(current_paragraph)
                    indent_px = 0
                if in_list:
                    html_parts.append(f'<li style="margin-bottom: 0.5rem; padding-left: 2rem; text-indent: -2rem;">{para_text}</li>')
                else:
                    if indent_px > 0:
                        html_parts.append(f'<p style="margin-left: {indent_px}px;">{para_text}</p>')
                    else:
                        html_parts.append(f'<p>{para_text}</p>')
                current_paragraph = []

        def flush_penjelasan_block():
            nonlocal in_penjelasan_block, penjelasan_lines
            if not in_penjelasan_block:
                return
            # First line is header, rest is content
            header = penjelasan_lines[0].strip() if penjelasan_lines else ''
            content = ' '.join([ln.strip() for ln in penjelasan_lines[1:]]).strip()
            # Render styled block; keep simple structure
            html_parts.append(
                '<div class="penjelasan" style="margin-top: 0.5rem; margin-bottom: 0.75rem; padding: 8px 12px; '
                'background: #f4f8ff; border-left: 4px solid #0d6efd; border-radius: 2px;">'
                f'<div style="font-weight: 600; color: #0d6efd;">{header}</div>'
                f'{(f"<div style=\"margin-top: 4px;\">{content}</div>" if content else "")}'
                '</div>'
            )
            penjelasan_lines = []
            in_penjelasan_block = False

        for line in lines:
            # PRESERVE LEADING SPACES for layout preservation
            # Count leading spaces before stripping
            leading_spaces = len(line) - len(line.lstrip())
            stripped = line.strip()
            
            # Handle penjelasan block rendering
            if in_penjelasan_block:
                # End of block on empty line
                if not stripped:
                    flush_penjelasan_block()
                    # continue to next line without treating this as paragraph break
                    continue
                else:
                    penjelasan_lines.append(line)
                    continue

            # Start of a penjelasan block
            if penjelasan_header_pat.match(stripped):
                # Flush any open paragraph or list before starting
                flush_current_paragraph()
                if in_list:
                    html_parts.append('</ul>')
                    in_list = False
                in_penjelasan_block = True
                penjelasan_lines = [line]
                continue

            # Preserve empty lines as paragraph breaks
            if not stripped:
                # If we were inside penjelasan (should have been handled above), flush
                if in_penjelasan_block:
                    flush_penjelasan_block()
                    continue
                if current_paragraph:
                    # Extract text and indentation from tuples
                    if isinstance(current_paragraph[0], tuple):
                        para_text = ' '.join([item[0] if isinstance(item, tuple) else item for item in current_paragraph])
                        indent_px = current_paragraph[0][1] if isinstance(current_paragraph[0], tuple) else 0
                    else:
                        para_text = ' '.join(current_paragraph)
                        indent_px = 0
                    
                    if in_list:
                        # For list continuation, don't add margin (controlled by <ul>)
                        html_parts.append(f'<li style="margin-bottom: 0.5rem; padding-left: 2rem; text-indent: -2rem;">{para_text}</li>')
                    else:
                        # Only add margin if there's actual indentation
                        if indent_px > 0:
                            html_parts.append(f'<p style="margin-left: {indent_px}px;">{para_text}</p>')
                        else:
                            html_parts.append(f'<p>{para_text}</p>')
                    current_paragraph = []
                elif in_list:
                    # Empty line while in list - close the list
                    html_parts.append('</ul>')
                    in_list = False
                continue
            
            # Check for legal document intro patterns (Mengingat:, Menimbang:, etc.)
            legal_intro_match = legal_intro_pattern.match(stripped)
            if legal_intro_match:
                if current_paragraph:
                    # Extract text and indentation
                    if isinstance(current_paragraph[0], tuple):
                        para_text = ' '.join([item[0] if isinstance(item, tuple) else item for item in current_paragraph])
                        indent_px = current_paragraph[0][1] if isinstance(current_paragraph[0], tuple) else 0
                    else:
                        para_text = ' '.join(current_paragraph)
                        indent_px = 0
                    
                    if in_list:
                        html_parts.append(f'<li style="margin-bottom: 0.5rem; padding-left: 2rem; text-indent: -2rem;">{para_text}</li>')
                    else:
                        # Only add margin if there's actual indentation
                        if indent_px > 0:
                            html_parts.append(f'<p style="margin-left: {indent_px}px;">{para_text}</p>')
                        else:
                            html_parts.append(f'<p>{para_text}</p>')
                    current_paragraph = []
                
                if in_list:
                    html_parts.append('</ul>')
                    in_list = False
                
                intro_word = legal_intro_match.group(1)
                rest = legal_intro_match.group(2).strip()
                
                # Add the intro as a heading
                html_parts.append(f'<h6 class="mt-3 mb-2"><strong>{intro_word}:</strong></h6>')
                
                # Check if there's numbering right after
                if rest and numbering_pattern.match(rest):
                    number_match = numbering_pattern.match(rest)
                    if number_match:
                        # Start new list
                        html_parts.append('<ul style="list-style-type: none; padding-left: 0; margin-left: 1.5rem;">')
                        in_list = True
                        number = number_match.group(1)
                        content = number_match.group(2)
                        # Use table display for perfect alignment
                        html_parts.append(f'<li style="margin-bottom: 0.5rem; display: table; width: 100%;"><span style="display: table-cell; width: 2rem; font-weight: bold; vertical-align: top;">{number}</span><span style="display: table-cell; word-wrap: break-word;">{content}</span></li>')
                elif rest:
                    html_parts.append(f'<p>{rest}</p>')
                
                continue
            
            # Check for section headers (BAB, PASAL, etc.)
            section_match = section_pattern.match(stripped)
            if section_match:
                if current_paragraph:
                    # Extract text and indentation
                    if isinstance(current_paragraph[0], tuple):
                        para_text = ' '.join([item[0] if isinstance(item, tuple) else item for item in current_paragraph])
                        indent_px = current_paragraph[0][1] if isinstance(current_paragraph[0], tuple) else 0
                    else:
                        para_text = ' '.join(current_paragraph)
                        indent_px = 0
                    
                    if in_list:
                        html_parts.append(f'<li style="margin-bottom: 0.5rem; padding-left: 2rem; text-indent: -2rem;">{para_text}</li>')
                    else:
                        # Only add margin if there's actual indentation
                        if indent_px > 0:
                            html_parts.append(f'<p style="margin-left: {indent_px}px;">{para_text}</p>')
                        else:
                            html_parts.append(f'<p>{para_text}</p>')
                    current_paragraph = []
                
                if in_list:
                    html_parts.append('</ul>')
                    in_list = False
                
                html_parts.append(f'<h5 class="mt-3 mb-2"><strong>{stripped}</strong></h5>')
                continue
            
            # Check for ALL CAPS header
            if len(stripped) > 10 and header_pattern.match(stripped):
                if current_paragraph:
                    # Extract text and indentation
                    if isinstance(current_paragraph[0], tuple):
                        para_text = ' '.join([item[0] if isinstance(item, tuple) else item for item in current_paragraph])
                        indent_px = current_paragraph[0][1] if isinstance(current_paragraph[0], tuple) else 0
                    else:
                        para_text = ' '.join(current_paragraph)
                        indent_px = 0
                    
                    if in_list:
                        html_parts.append(f'<li style="margin-bottom: 0.5rem; padding-left: 2rem; text-indent: -2rem;">{para_text}</li>')
                    else:
                        # Only add margin if there's actual indentation
                        if indent_px > 0:
                            html_parts.append(f'<p style="margin-left: {indent_px}px;">{para_text}</p>')
                        else:
                            html_parts.append(f'<p>{para_text}</p>')
                    current_paragraph = []
                
                if in_list:
                    html_parts.append('</ul>')
                    in_list = False
                
                html_parts.append(f'<h6 class="mt-2 mb-2"><strong>{stripped}</strong></h6>')
                continue
            
            # Check for numbered/lettered lists - PRESERVE ORIGINAL FORMAT
            numbering_match = numbering_pattern.match(stripped)
            if numbering_match:
                if current_paragraph:
                    # Extract text and indentation
                    if isinstance(current_paragraph[0], tuple):
                        para_text = ' '.join([item[0] if isinstance(item, tuple) else item for item in current_paragraph])
                        para_indent_px = current_paragraph[0][1] if isinstance(current_paragraph[0], tuple) else 0
                    else:
                        para_text = ' '.join(current_paragraph)
                        para_indent_px = 0
                    
                    if in_list:
                        # Apply hanging indent to multi-line content
                        html_parts.append(f'<li style="margin-bottom: 0.5rem; padding-left: 2rem; text-indent: -2rem;">{para_text}</li>')
                    else:
                        html_parts.append(f'<p style="margin-left: {para_indent_px}px;">{para_text}</p>')
                    current_paragraph = []
                
                if not in_list:
                    # Start new list with fixed left margin (no nested indentation)
                    html_parts.append('<ul style="list-style-type: none; padding-left: 0; margin-left: 1.5rem;">')
                    in_list = True
                
                # Add list item with ORIGINAL numbering and hanging indent
                # Use table display for perfect alignment
                # For list items, DON'T add margin-left (already controlled by <ul>)
                # Just use table display for proper alignment
                number = numbering_match.group(1)
                content = numbering_match.group(2)
                html_parts.append(f'<li style="margin-bottom: 0.5rem; display: table; width: 100%;"><span style="display: table-cell; width: 2rem; font-weight: bold; vertical-align: top;">{number}</span><span style="display: table-cell; word-wrap: break-word;">{content}</span></li>')
                continue
            
            # If not a special pattern, accumulate into current paragraph
            # This preserves line breaks within sections
            # Store with indentation info
            if leading_spaces > 0:
                indent_px = leading_spaces * 8
                current_paragraph.append((stripped, indent_px))
            else:
                current_paragraph.append((stripped, 0))
        
        # Close any remaining content
        # Flush penjelasan block if still open
        if in_penjelasan_block:
            flush_penjelasan_block()
        if current_paragraph:
            # Extract text and indentation
            if isinstance(current_paragraph[0], tuple):
                para_text = ' '.join([item[0] if isinstance(item, tuple) else item for item in current_paragraph])
                indent_px = current_paragraph[0][1] if isinstance(current_paragraph[0], tuple) else 0
            else:
                para_text = ' '.join(current_paragraph)
                indent_px = 0
            
            if in_list:
                # Apply hanging indent to closing content
                html_parts.append(f'<li style="margin-bottom: 0.5rem; padding-left: 2rem; text-indent: -2rem;">{para_text}</li>')
                html_parts.append('</ul>')
            else:
                # Only add margin if there's actual indentation
                if indent_px > 0:
                    html_parts.append(f'<p style="margin-left: {indent_px}px;">{para_text}</p>')
                else:
                    html_parts.append(f'<p>{para_text}</p>')
        elif in_list:
            html_parts.append('</ul>')
        
        return ''.join(html_parts)
    
    def _extract_text_from_pdf(self, pdf_data):
        """Extract text content from PDF file with improved formatting - uses advanced PDFExtractor"""
        try:
            # Import the advanced PDF parser
            from .parser import PDFExtractor
            
            # Decode base64 PDF data
            pdf_bytes = base64.b64decode(pdf_data)
            
            # Use the advanced PDFExtractor
            _logger.info("Using advanced PDFExtractor for PDF text extraction...")
            extractor = PDFExtractor(pdf_bytes)
            
            # First try layout-preserving extraction
            _logger.info("=" * 80)
            _logger.info("LAYOUT-PRESERVING EXTRACTION STARTED")
            _logger.info("=" * 80)
            text = extractor.extract_text_with_layout()
            
            if text.strip():
                _logger.info(f"✅ Layout extraction SUCCESS: {len(text)} characters")
                _logger.info("First 500 chars of layout extraction:")
                _logger.info(text[:500])
            else:
                _logger.warning("⚠️ Layout extraction returned empty, falling back to auto extraction...")
                text = extractor.extract_text_auto()
                _logger.info(f"Auto extraction: {len(text)} characters")
            
            if text.strip():
                _logger.info(f"PDFExtractor successfully extracted {len(text)} characters")
                
                # Format the extracted text with our HTML formatter
                formatted_text = self._format_text_to_html(text)
                
                result = f'<div class="pdf-content" style="line-height: 1.8;">{formatted_text}</div>'
                return result
            else:
                _logger.warning("PDFExtractor returned empty text")
                return '<p class="text-warning"><strong>Perhatian:</strong> Tidak dapat mengekstrak teks dari PDF ini. PDF mungkin berisi gambar atau format yang tidak didukung.</p>'
                
        except ImportError as ie:
            # Fallback to old method if parser.py or its dependencies not available
            _logger.warning(f"PDFExtractor not available ({ie}), falling back to basic extraction...")
            
            try:
                pdf_bytes = base64.b64decode(pdf_data)
                pdf_file = io.BytesIO(pdf_bytes)
                
                # Try PyPDF2 first
                try:
                    from PyPDF2 import PdfReader
                    reader = PdfReader(pdf_file)
                    html_content = ['<div class="pdf-content" style="line-height: 1.8;">']
                    
                    for page_num, page in enumerate(reader.pages):
                        try:
                            text = page.extract_text()
                            if text.strip():
                                html_content.append(f'<div class="page-break mb-4">')
                                html_content.append(f'<h4 class="text-primary border-bottom pb-2 mb-3">Halaman {page_num + 1}</h4>')
                                
                                # Format text with proper structure
                                formatted_text = self._format_text_to_html(text)
                                html_content.append(formatted_text)
                                html_content.append('</div>')
                        except Exception as e:
                            _logger.warning(f"Error extracting page {page_num + 1}: {e}")
                            continue
                    
                    html_content.append('</div>')
                    
                    if len(html_content) > 2:  # More than just opening and closing div
                        return ''.join(html_content)
                        
                except ImportError:
                    _logger.warning("PyPDF2 not installed, trying pdfminer.six...")
                    
                # Fallback to pdfminer.six
                try:
                    from pdfminer.high_level import extract_text
                    pdf_file.seek(0)  # Reset file pointer
                    text = extract_text(pdf_file)
                    
                    if text.strip():
                        # Format text with proper structure
                        formatted_text = self._format_text_to_html(text)
                        return f'<div class="pdf-content" style="line-height: 1.8;">{formatted_text}</div>'
                        
                except ImportError:
                    _logger.error("Neither PyPDF2 nor pdfminer.six is installed")
                    return '<p class="text-danger"><strong>Error:</strong> Library untuk ekstraksi PDF tidak tersedia. Silakan install PyPDF2 atau pdfminer.six</p>'
                
                return '<p class="text-warning"><strong>Perhatian:</strong> Tidak dapat mengekstrak teks dari PDF ini.</p>'
                
            except Exception as fallback_error:
                _logger.error(f"Fallback extraction failed: {fallback_error}")
                return f'<p class="text-danger"><strong>Error:</strong> Gagal mengekstrak teks dari PDF: {fallback_error}</p>'
            
        except Exception as e:
            _logger.error(f"Error extracting PDF text: {str(e)}", exc_info=True)
            return f'<p class="text-danger"><strong>Error:</strong> Gagal mengekstrak teks dari PDF: {str(e)}</p>'
    
    def _get_paragraph_xml_text(self, paragraph):
        """Get raw XML text from paragraph including numbering"""
        try:
            from lxml import etree
            
            # Get the paragraph's XML element
            p_element = paragraph._element
            
            # Try to get any numbering or list information from XML
            xml_str = etree.tostring(p_element, encoding='unicode', method='text')
            
            return xml_str.strip() if xml_str else None
            
        except Exception as e:
            _logger.debug(f"Could not get XML text: {str(e)}")
            return None
    
    def _detect_list_format(self, text):
        """Detect if text starts with list numbering and extract it"""
        import re
        
        if not text:
            return None, text
        
        # Patterns untuk berbagai format numbering
        patterns = [
            r'^([a-z]\.)\s+(.+)$',  # a. text
            r'^([a-z]\))\s+(.+)$',  # a) text
            r'^(\d+\.)\s+(.+)$',    # 1. text
            r'^(\d+\))\s+(.+)$',    # 1) text
            r'^([ivxlc]+\.)\s+(.+)$',  # i. text (roman)
            r'^(\([a-z]\))\s+(.+)$',  # (a) text
            r'^(\(\d+\))\s+(.+)$',   # (1) text
        ]
        
        for pattern in patterns:
            match = re.match(pattern, text.strip(), re.IGNORECASE)
            if match:
                return match.group(1), match.group(2)
        
        return None, text
    
    def _extract_text_from_docx(self, docx_data):
        """Extract text from DOCX file - MUCH MORE ACCURATE than PDF!"""
        try:
            from docx import Document
            from lxml import etree
            import re
            
            # Decode base64 DOCX data
            docx_bytes = base64.b64decode(docx_data)
            docx_file = io.BytesIO(docx_bytes)
            
            doc = Document(docx_file)
            
            # Extract all paragraphs with proper structure
            all_text = []
            list_item_counter = {}  # Track list items per numbering ID
            
            for para in doc.paragraphs:
                text = para.text.strip()
                
                if not text:
                    # Empty paragraph = line break
                    all_text.append('')
                    continue
                
                # Check if this paragraph has Word numbering format
                has_word_numbering = False
                generated_numbering = None
                
                try:
                    if para._element.pPr is not None:
                        numPr = para._element.pPr.numPr
                        if numPr is not None:
                            # This paragraph has numbering format in Word
                            has_word_numbering = True
                            
                            # Try to get numbering level
                            ilvl = 0
                            numId = 0
                            if numPr.ilvl is not None:
                                ilvl = int(numPr.ilvl.val)
                            if numPr.numId is not None:
                                numId = int(numPr.numId.val)
                            
                            # Track counter for this numbering ID and level
                            counter_key = f"{numId}_{ilvl}"
                            if counter_key not in list_item_counter:
                                list_item_counter[counter_key] = 0
                            list_item_counter[counter_key] += 1
                            count = list_item_counter[counter_key]
                            
                            # Generate numbering based on level
                            if ilvl == 0:
                                # Lowercase letters for level 0
                                if count <= 26:
                                    generated_numbering = f"{chr(96 + count)}."
                                else:
                                    generated_numbering = f"{count}."
                            elif ilvl == 1:
                                # Numbers for level 1
                                generated_numbering = f"{count}."
                            else:
                                # Default to numbers
                                generated_numbering = f"{count}."
                            
                except Exception as e:
                    _logger.debug(f"Numbering detection from structure failed: {str(e)}")
                
                # Now decide what to do based on what we found
                if has_word_numbering and generated_numbering:
                    # This paragraph has Word numbering
                    # Check if the text ALREADY starts with the same or similar numbering
                    detected_num, remaining_text = self._detect_list_format(text)
                    
                    if detected_num:
                        # Text already has numbering - DON'T add generated numbering
                        all_text.append(text)
                        _logger.info(f"Paragraph has Word numbering BUT text already has it: {text[:50]}...")
                    else:
                        # Text doesn't have numbering yet - ADD generated numbering
                        full_text = f"{generated_numbering} {text}"
                        all_text.append(full_text)
                        _logger.info(f"Added generated numbering {generated_numbering} to: {text[:30]}...")
                else:
                    # No Word numbering detected
                    # Just use the text as-is (may or may not have manual numbering)
                    all_text.append(text)
            
            # IMPORTANT: Keep line breaks when joining
            # This preserves paragraph structure and list items
            full_text = '\n'.join(all_text)
            
            _logger.info(f"DOCX extraction: {len(all_text)} paragraphs extracted")
            _logger.info(f"First 500 chars: {full_text[:500]}")
            
            # Format using our smart formatter (will detect Menimbang:, numbering, etc.)
            formatted_text = self._format_text_to_html(full_text)
            
            result = f'<div class="docx-content" style="line-height: 1.8;">{formatted_text}</div>'
            
            if len(result) > 50:  # Has content
                _logger.info(f"DOCX extraction successful: {len(all_text)} paragraphs, {len(full_text)} chars")
                return result
            else:
                return '<p class="text-warning"><strong>Perhatian:</strong> File DOCX kosong atau tidak dapat dibaca.</p>'
                
        except ImportError as ie:
            missing_lib = str(ie).split("'")[1] if "'" in str(ie) else "unknown"
            _logger.error(f"Missing library: {missing_lib}")
            return f'<p class="text-danger"><strong>Error:</strong> Library {missing_lib} tidak terinstall.</p>'
        except Exception as e:
            _logger.error(f"Error extracting DOCX text: {str(e)}", exc_info=True)
            return f'<p class="text-danger"><strong>Error:</strong> Gagal mengekstrak teks dari DOCX: {str(e)}</p>'
    
    def _convert_docx_to_pdf(self, docx_data):
        """Convert DOCX to PDF using python-docx and reportlab"""
        try:
            from docx import Document
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            
            # Decode DOCX
            docx_bytes = base64.b64decode(docx_data)
            docx_file = io.BytesIO(docx_bytes)
            doc = Document(docx_file)
            
            # Create PDF in memory
            pdf_buffer = io.BytesIO()
            pdf_doc = SimpleDocTemplate(pdf_buffer, pagesize=A4,
                                       rightMargin=72, leftMargin=72,
                                       topMargin=72, bottomMargin=18)
            
            # Container for PDF elements
            elements = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            style_normal = ParagraphStyle('CustomNormal', parent=styles['Normal'],
                                         fontSize=11, leading=14, alignment=TA_JUSTIFY)
            style_heading = ParagraphStyle('CustomHeading', parent=styles['Heading1'],
                                          fontSize=14, leading=16, spaceAfter=12)
            
            # Extract and add content
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    elements.append(Spacer(1, 0.2*inch))
                    continue
                
                # Escape HTML special chars
                text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                
                # Check if heading (ALL CAPS or starts with BAB, PASAL, etc)
                if text.isupper() and len(text) > 10:
                    elements.append(Paragraph(f"<b>{text}</b>", style_heading))
                elif text.upper().startswith(('BAB ', 'PASAL ', 'AYAT ', 'BAGIAN ')):
                    elements.append(Paragraph(f"<b>{text}</b>", style_heading))
                else:
                    elements.append(Paragraph(text, style_normal))
            
            # Build PDF
            pdf_doc.build(elements)
            
            # Get PDF bytes and encode to base64
            pdf_bytes = pdf_buffer.getvalue()
            pdf_base64 = base64.b64encode(pdf_bytes)
            
            return pdf_base64
            
        except ImportError as e:
            _logger.error(f"Missing library for PDF conversion: {e}. Install: pip install python-docx reportlab")
            return None
        except Exception as e:
            _logger.error(f"Error converting DOCX to PDF: {str(e)}")
            return None
    
    def _parse_pasal_structure(self, text):
        """Parse text into Pasal-Ayat-Huruf structure with Penjelasan support"""
        lines = text.splitlines()
        pasal_blocks = []
        toc = []
        
        current_pasal = None
        current_pasal_data = None
        current_ayat = None
        in_penjelasan_section = False
        penjelasan_ayat_map = {}
        
        # Regex patterns
        pasal_pat = re.compile(r'^Pasal\s+(\d+)')
        ayat_pat = re.compile(r'^\((\d+)\)\s*(.*)')
        huruf_pat = re.compile(r'^\s{4}([a-z])\.\s*(.*)')
        penjelasan_section_pat = re.compile(r'^Penjelasan\s*$')
        penjelasan_pasal_pat = re.compile(r'^Pasal\s+(\d+)\s*$')
        penjelasan_ayat_pat = re.compile(r'^Penjelasan\s+\((\d+)\):\s*(.*)')
        penjelasan_huruf_pat = re.compile(r'^Penjelasan\s+\((\d+)\)([a-z]):\s*(.*)')
        bab_pat = re.compile(r'^BAB\s+[IVX]+')
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue
            
            if bab_pat.match(line):
                i += 1
                continue
            
            if penjelasan_section_pat.match(line):
                in_penjelasan_section = True
                i += 1
                continue
            
            if in_penjelasan_section:
                pasal_penjelasan_match = penjelasan_pasal_pat.match(line)
                if pasal_penjelasan_match:
                    current_pasal_penjelasan = pasal_penjelasan_match.group(1)
                    i += 1
                    temp_penjelasan = []
                    while i < len(lines):
                        next_line = lines[i].strip()
                        if not next_line:
                            i += 1
                            continue
                        if penjelasan_ayat_pat.match(next_line) or penjelasan_pasal_pat.match(next_line) or pasal_pat.match(next_line):
                            break
                        temp_penjelasan.append(next_line)
                        i += 1
                    
                    if temp_penjelasan:
                        penjelasan_text = ' '.join(temp_penjelasan)
                        if current_pasal_penjelasan not in penjelasan_ayat_map:
                            penjelasan_ayat_map[current_pasal_penjelasan] = {'umum': penjelasan_text, 'ayat': {}}
                        else:
                            penjelasan_ayat_map[current_pasal_penjelasan]['umum'] = penjelasan_text
                    continue
                
                penjelasan_ayat_match = penjelasan_ayat_pat.match(line)
                if penjelasan_ayat_match and current_pasal:
                    ayat_num = penjelasan_ayat_match.group(1)
                    penjelasan_text = penjelasan_ayat_match.group(2).strip()
                    
                    if current_pasal not in penjelasan_ayat_map:
                        penjelasan_ayat_map[current_pasal] = {'umum': '', 'ayat': {}}
                    penjelasan_ayat_map[current_pasal]['ayat'][ayat_num] = penjelasan_text
                    i += 1
                    continue
                
                penjelasan_huruf_match = penjelasan_huruf_pat.match(line)
                if penjelasan_huruf_match and current_pasal:
                    ayat_num = penjelasan_huruf_match.group(1)
                    huruf = penjelasan_huruf_match.group(2)
                    penjelasan_text = penjelasan_huruf_match.group(3).strip()
                    
                    if current_pasal not in penjelasan_ayat_map:
                        penjelasan_ayat_map[current_pasal] = {'umum': '', 'ayat': {}}
                    if ayat_num not in penjelasan_ayat_map[current_pasal]['ayat']:
                        penjelasan_ayat_map[current_pasal]['ayat'][ayat_num] = {}
                    if not isinstance(penjelasan_ayat_map[current_pasal]['ayat'][ayat_num], dict):
                        penjelasan_ayat_map[current_pasal]['ayat'][ayat_num] = {'_text': penjelasan_ayat_map[current_pasal]['ayat'][ayat_num]}
                    
                    penjelasan_ayat_map[current_pasal]['ayat'][ayat_num][huruf] = penjelasan_text
                    i += 1
                    continue
            
            pasal_match = pasal_pat.match(line)
            if pasal_match and not in_penjelasan_section:
                if current_pasal_data:
                    pasal_blocks.append(current_pasal_data)
                    toc.append({'pasal': current_pasal, 'ayats': [a['nomor'] for a in current_pasal_data['ayats']]})
                
                current_pasal = pasal_match.group(1)
                current_pasal_data = {'nomor': current_pasal, 'ayats': [], 'penjelasan_umum': ''}
                current_ayat = None
                i += 1
                continue
            
            ayat_match = ayat_pat.match(line)
            if ayat_match and current_pasal_data and not in_penjelasan_section:
                ayat_num = ayat_match.group(1)
                ayat_text = ayat_match.group(2).strip()
                
                current_ayat = {'nomor': ayat_num, 'isi': [ayat_text] if ayat_text else [], 'hurufs': [], 'penjelasan': ''}
                current_pasal_data['ayats'].append(current_ayat)
                i += 1
                continue
            
            huruf_match = huruf_pat.match(line)
            if huruf_match and current_ayat and not in_penjelasan_section:
                huruf = huruf_match.group(1)
                huruf_text = huruf_match.group(2).strip()
                current_ayat['hurufs'].append({'huruf': huruf, 'isi': huruf_text, 'penjelasan': ''})
                i += 1
                continue
            
            if current_ayat and not in_penjelasan_section and line:
                current_ayat['isi'].append(line)
            
            i += 1
        
        if current_pasal_data:
            pasal_blocks.append(current_pasal_data)
            toc.append({'pasal': current_pasal, 'ayats': [a['nomor'] for a in current_pasal_data['ayats']]})
        
        # Merge penjelasan
        for pasal in pasal_blocks:
            pasal_num = pasal['nomor']
            if pasal_num in penjelasan_ayat_map:
                pasal['penjelasan_umum'] = penjelasan_ayat_map[pasal_num].get('umum', '')
                
                for ayat in pasal['ayats']:
                    ayat_num = ayat['nomor']
                    if ayat_num in penjelasan_ayat_map[pasal_num].get('ayat', {}):
                        penjelasan_data = penjelasan_ayat_map[pasal_num]['ayat'][ayat_num]
                        
                        if isinstance(penjelasan_data, str):
                            ayat['penjelasan'] = penjelasan_data
                        elif isinstance(penjelasan_data, dict):
                            ayat['penjelasan'] = penjelasan_data.get('_text', '')
                            
                            for huruf_item in ayat['hurufs']:
                                huruf = huruf_item['huruf']
                                if huruf in penjelasan_data:
                                    huruf_item['penjelasan'] = penjelasan_data[huruf]
        
        return pasal_blocks, toc

    def _generate_structured_html(self, pasal_blocks, toc):
        """Generate HTML with accordion + TOC + Penjelasan"""
        html_parts = []
        
        # CSS
        html_parts.append('''
<style>
.regulation-container { display: flex; gap: 20px; }
.toc-sidebar { flex: 0 0 250px; position: sticky; top: 20px; max-height: 80vh; overflow-y: auto; 
               padding: 15px; background: #f8f9fa; border-radius: 8px; border: 1px solid #dee2e6; }
.toc-sidebar h5 { margin-bottom: 15px; color: #495057; font-weight: 600; font-size: 1em; }
.toc-sidebar ul { list-style: none; padding-left: 0; margin: 0; }
.toc-sidebar ul ul { padding-left: 15px; margin-top: 5px; }
.toc-sidebar a { color: #007bff; text-decoration: none; display: block; padding: 3px 0; font-size: 0.9em; }
.toc-sidebar a:hover { color: #0056b3; text-decoration: underline; }
.regulation-content { flex: 1; }
.pasal-card { margin-bottom: 20px; border: 1px solid #dee2e6; border-radius: 8px; overflow: hidden; }
.pasal-header { background: #007bff; color: white; padding: 12px 15px; cursor: pointer; 
                font-weight: 600; font-size: 1.05em; }
.pasal-header:hover { background: #0056b3; }
.pasal-content { padding: 20px; background: white; }
.ayat { margin-bottom: 15px; padding: 12px; background: #f8f9fa; border-left: 4px solid #007bff; border-radius: 4px; }
.ayat strong { color: #007bff; font-size: 1.05em; }
.ayat ul { margin-top: 8px; padding-left: 20px; list-style: none; }
.ayat li { margin-bottom: 8px; }
.penjelasan { margin-top: 10px; padding: 10px; background: #fff3cd; border-left: 4px solid #ffc107; 
              border-radius: 4px; font-style: italic; font-size: 0.95em; }
.penjelasan strong { color: #856404; font-style: normal; }
.penjelasan-umum { margin-bottom: 15px; padding: 12px; background: #d1ecf1; border-left: 4px solid #17a2b8; 
                   border-radius: 4px; font-size: 0.95em; }
.penjelasan-umum strong { color: #0c5460; }
</style>
''')
        
        html_parts.append('<div class="regulation-container">')
        
        # TOC Sidebar
        html_parts.append('<div class="toc-sidebar"><h5>📋 Daftar Isi</h5><ul>')
        for entry in toc:
            html_parts.append(f'<li><a href="#pasal{entry["pasal"]}">Pasal {entry["pasal"]}</a>')
            if entry['ayats']:
                html_parts.append('<ul>')
                for ayat in entry['ayats']:
                    html_parts.append(f'<li><a href="#pasal{entry["pasal"]}-ayat{ayat}">Ayat ({ayat})</a></li>')
                html_parts.append('</ul>')
            html_parts.append('</li>')
        html_parts.append('</ul></div>')
        
        # Content
        html_parts.append('<div class="regulation-content">')
        
        for pasal in pasal_blocks:
            pasal_id = f'pasal{pasal["nomor"]}'
            html_parts.append(f'<div class="pasal-card" id="{pasal_id}">')
            html_parts.append(f'<div class="pasal-header" onclick="this.nextElementSibling.style.display=this.nextElementSibling.style.display===\'none\'?\'block\':\'none\'">📜 Pasal {pasal["nomor"]}</div>')
            html_parts.append(f'<div class="pasal-content" style="display:block;">')
            
            if pasal.get('penjelasan_umum'):
                html_parts.append(f'<div class="penjelasan-umum"><strong>💡 Penjelasan Pasal:</strong><br>{pasal["penjelasan_umum"]}</div>')
            
            for ayat in pasal['ayats']:
                ayat_id = f'{pasal_id}-ayat{ayat["nomor"]}'
                ayat_text = ' '.join(ayat['isi']).strip()
                html_parts.append(f'<div class="ayat" id="{ayat_id}"><strong>({ayat["nomor"]})</strong> {ayat_text}')
                
                if ayat['hurufs']:
                    html_parts.append('<ul>')
                    for huruf in ayat['hurufs']:
                        html_parts.append(f'<li><strong>{huruf["huruf"]}.</strong> {huruf["isi"]}')
                        if huruf.get('penjelasan'):
                            html_parts.append(f'<div class="penjelasan"><strong>💡 Penjelasan huruf {huruf["huruf"]}:</strong> {huruf["penjelasan"]}</div>')
                        html_parts.append('</li>')
                    html_parts.append('</ul>')
                
                if ayat.get('penjelasan'):
                    html_parts.append(f'<div class="penjelasan"><strong>💡 Penjelasan ayat ({ayat["nomor"]}):</strong> {ayat["penjelasan"]}</div>')
                
                html_parts.append('</div>')
            
            html_parts.append('</div></div>')
        
        html_parts.append('</div></div>')
        
        return '\n'.join(html_parts)

    def _insert_penjelasan_into_text(self, text):
        """
        Insert Penjelasan sections into their respective Pasal/Ayat/Huruf positions.
        
        Handles complex hierarchical structure:
        - KUBU 1 (Main Content): Regular regulation text with Pasal, Ayat, Huruf
        - KUBU 2 (Penjelasan): Starts with "PENJELASAN ATAS" or "Penjelasan"
        
        Format in Penjelasan section:
        - Pasal X (standalone) = explanation for whole Pasal
        - Ayat (X) (under context of Pasal) = explanation for specific Ayat
        - Huruf X (under context of Pasal and Ayat) = explanation for specific Huruf
        """
        try:
            lines = text.split('\n')
            result_lines = []
            penjelasan_map = {}  # key: "pasal_X", "pasal_X_ayat_Y", "pasal_X_ayat_Y_huruf_Z"
            in_penjelasan_section = False
            current_penjelasan_context = {'pasal': None, 'ayat': None}
            current_penjelasan_key = None
            current_penjelasan_buffer = []
            
            # Regex patterns for detecting Penjelasan section
            penjelasan_section_pat = re.compile(r'^(PENJELASAN\s+ATAS|Penjelasan)\s*$', re.IGNORECASE)
            
            # Patterns in Penjelasan section
            pasal_in_penjelasan_pat = re.compile(r'^Pasal\s+(\d+)\s*$', re.IGNORECASE)
            # FIXED: Support both "Ayat (1)" in main content and "Ayat 1" in Penjelasan section
            ayat_header_pat = re.compile(r'^Ayat\s+\(?(\d+)\)?\s*$', re.IGNORECASE)
            # FIXED: Match "Huruf x" exactly (standalone, not inline with text)
            huruf_header_pat = re.compile(r'^Huruf\s+([a-z])\s*$', re.IGNORECASE)
            
            # Pattern to detect "Cukup jelas" (means no explanation needed)
            cukup_jelas_pat = re.compile(r'^Cukup\s+jelas\.?\s*$', re.IGNORECASE)
            
            # Pattern to detect continuation markers (e.g., "Huruf d ...", "Pasal 2 ...")
            continuation_marker_pat = re.compile(r'^(Pasal|Ayat|Huruf)\s+.*\.\.\.\s*$', re.IGNORECASE)
            
            _logger.info("=" * 80)
            _logger.info("PARSING PENJELASAN SECTION")
            _logger.info("=" * 80)
            
            # FIRST PASS: Collect all penjelasan entries
            i = 0
            while i < len(lines):
                line = lines[i].strip()
                
                # Detect start of Penjelasan section
                if penjelasan_section_pat.match(line):
                    in_penjelasan_section = True
                    _logger.info(f"✓ Found Penjelasan section at line {i}: {line}")
                    i += 1
                    continue
                
                if in_penjelasan_section:
                    # Check for Pasal declaration (standalone = explanation for whole Pasal)
                    pasal_match = pasal_in_penjelasan_pat.match(line)
                    if pasal_match:
                        # Save previous penjelasan if any
                        if current_penjelasan_key and current_penjelasan_buffer:
                            content = ' '.join(current_penjelasan_buffer).strip()
                            if content and not cukup_jelas_pat.match(content):
                                penjelasan_map[current_penjelasan_key] = content
                                _logger.info(f"  Saved [{current_penjelasan_key}]: {content[:50]}...")
                        
                        # Set new context
                        pasal_num = pasal_match.group(1)
                        current_penjelasan_context = {'pasal': pasal_num, 'ayat': None}
                        
                        # Peek ahead: check if next non-empty line is "Ayat" or "Huruf"
                        # If yes, don't set penjelasan_key to pasal (no pasal-level explanation)
                        peek_ahead_idx = i + 1
                        has_immediate_ayat_or_huruf = False
                        while peek_ahead_idx < len(lines):
                            peek_line = lines[peek_ahead_idx].strip()
                            if peek_line:
                                # Check if it's an Ayat or Huruf header
                                if (ayat_header_pat.match(peek_line) or 
                                    huruf_header_pat.match(peek_line)):
                                    has_immediate_ayat_or_huruf = True
                                break
                            peek_ahead_idx += 1
                        
                        if has_immediate_ayat_or_huruf:
                            # Don't set pasal key - pasal has no direct explanation
                            current_penjelasan_key = None
                            _logger.info(f"\n→ Pasal {pasal_num} context started (no pasal-level explanation, has Ayat/Huruf)")
                        else:
                            # Set pasal key - pasal has explanation
                            current_penjelasan_key = f'pasal_{pasal_num}'
                            _logger.info(f"\n→ Pasal {pasal_num} context started (has pasal-level explanation)")
                        
                        current_penjelasan_buffer = []
                        i += 1
                        continue
                    
                    # Check for Ayat header (e.g., "Ayat (1)")
                    ayat_match = ayat_header_pat.match(line)
                    if ayat_match and current_penjelasan_context['pasal']:
                        # Save previous penjelasan (could be for Pasal or previous Ayat)
                        if current_penjelasan_key and current_penjelasan_buffer:
                            content = ' '.join(current_penjelasan_buffer).strip()
                            # Check if content is meaningful (not just "Ayat X Cukup jelas" repeated)
                            # Filter out buffer that only contains headers like "Ayat 1", "Ayat 2", etc
                            filtered_content = ' '.join([
                                line for line in current_penjelasan_buffer 
                                if line.strip() and 
                                not re.match(r'^Ayat\s+\d+\s*$', line.strip()) and
                                not re.match(r'^Huruf\s+[a-z]\s*$', line.strip())
                            ]).strip()
                            
                            if filtered_content and len(filtered_content) > 10:
                                penjelasan_map[current_penjelasan_key] = filtered_content
                                _logger.info(f"  Saved [{current_penjelasan_key}]: {filtered_content[:50]}...")
                        
                        # Set new Ayat context
                        ayat_num = ayat_match.group(1)
                        current_penjelasan_context['ayat'] = ayat_num
                        pasal_num = current_penjelasan_context['pasal']
                        
                        # Peek ahead: check if next non-empty line is "Huruf"
                        peek_ahead_idx = i + 1
                        has_immediate_huruf = False
                        while peek_ahead_idx < len(lines):
                            peek_line = lines[peek_ahead_idx].strip()
                            if peek_line:
                                if huruf_header_pat.match(peek_line):
                                    has_immediate_huruf = True
                                break
                            peek_ahead_idx += 1
                        
                        if has_immediate_huruf:
                            # Don't set ayat key - ayat has no direct explanation, has Huruf
                            current_penjelasan_key = None
                            _logger.info(f"  → Ayat ({ayat_num}) context started (no ayat-level explanation, has Huruf)")
                        else:
                            # Set ayat key - ayat has explanation
                            current_penjelasan_key = f"pasal_{pasal_num}_ayat_{ayat_num}"
                            _logger.info(f"  → Ayat ({ayat_num}) context started (has ayat-level explanation)")
                        
                        current_penjelasan_buffer = []
                        i += 1
                        continue
                    
                    # Check for Huruf header (e.g., "Huruf a")
                    huruf_match = huruf_header_pat.match(line)
                    if huruf_match and current_penjelasan_context['pasal'] and current_penjelasan_context['ayat']:
                        # Save previous penjelasan
                        if current_penjelasan_key and current_penjelasan_buffer:
                            content = ' '.join(current_penjelasan_buffer).strip()
                            # Filter out headers
                            filtered_content = ' '.join([
                                line for line in current_penjelasan_buffer 
                                if line.strip() and 
                                not re.match(r'^Ayat\s+\d+\s*$', line.strip()) and
                                not re.match(r'^Huruf\s+[a-z]\s*$', line.strip())
                            ]).strip()
                            
                            if filtered_content and len(filtered_content) > 10:
                                penjelasan_map[current_penjelasan_key] = filtered_content
                                _logger.info(f"  Saved [{current_penjelasan_key}]: {filtered_content[:50]}...")
                        
                        # Set new Huruf context
                        huruf = huruf_match.group(1)
                        pasal_num = current_penjelasan_context['pasal']
                        ayat_num = current_penjelasan_context['ayat']
                        current_penjelasan_key = f"pasal_{pasal_num}_ayat_{ayat_num}_huruf_{huruf}"
                        current_penjelasan_buffer = []
                        _logger.info(f"    → Huruf {huruf} context started")
                        i += 1
                        continue
                    
                    # Accumulate text for current penjelasan (skip empty lines at start)
                    if line:
                        # Skip "Cukup jelas", page markers, continuation markers, and headers
                        if (not cukup_jelas_pat.match(line) and 
                            not continuation_marker_pat.match(line) and
                            not re.match(r'^\s*-\s*\d+\s*-\s*$', line) and  # Page markers
                            not re.match(r'^Ayat\s+\d+\s*$', line.strip()) and  # Ayat headers
                            not re.match(r'^Huruf\s+[a-z]\s*$', line.strip())):  # Huruf headers
                            if current_penjelasan_buffer or line.strip():  # Skip leading empty lines
                                current_penjelasan_buffer.append(line)
                
                i += 1
            
            # Save last penjelasan (only if content is not empty and not just "Cukup jelas")
            if current_penjelasan_key and current_penjelasan_buffer:
                content = ' '.join(current_penjelasan_buffer).strip()
                # Double check: content should have meaningful text
                if content and len(content) > 10:  # More than just whitespace or short placeholder
                    penjelasan_map[current_penjelasan_key] = content
                    _logger.info(f"  Saved [{current_penjelasan_key}]: {content[:50]}...")
            
            _logger.info(f"\n✓ Total penjelasan collected: {len(penjelasan_map)}")
            for key in sorted(penjelasan_map.keys()):
                _logger.info(f"  - {key}: {penjelasan_map[key][:60]}...")
            
            # SECOND PASS: Insert penjelasan into main content
            in_penjelasan_section = False
            current_pasal = None
            current_ayat = None
            
            # pending penjelasan markers to flush at boundaries
            pending_pasal_lines = None
            pending_ayat_lines = None
            pending_huruf_lines = None
            
            # Patterns for main content
            pasal_main_pat = re.compile(r'^Pasal\s+(\d+)', re.IGNORECASE)
            ayat_main_pat = re.compile(r'^\((\d+)\)')
            huruf_main_pat = re.compile(r'^([a-z])\.')
            bab_header_pat = re.compile(r'^(BAB|BAGIAN)\s+', re.IGNORECASE)
            
            _logger.info("\n" + "=" * 80)
            _logger.info("INSERTING PENJELASAN INTO MAIN CONTENT (AFTER CONTENT)")
            _logger.info("=" * 80)
            
            def flush_huruf():
                nonlocal pending_huruf_lines
                if pending_huruf_lines:
                    result_lines.extend(pending_huruf_lines)
                    pending_huruf_lines = None
            
            def flush_ayat():
                nonlocal pending_ayat_lines
                # Ensure huruf is flushed before ayat
                flush_huruf()
                if pending_ayat_lines:
                    result_lines.extend(pending_ayat_lines)
                    pending_ayat_lines = None
            
            def flush_pasal():
                nonlocal pending_pasal_lines
                # Ensure ayat (and huruf) are flushed before pasal
                flush_ayat()
                if pending_pasal_lines:
                    result_lines.extend(pending_pasal_lines)
                    pending_pasal_lines = None
            
            for line in lines:
                stripped = line.strip()
                
                # Stop processing when we hit Penjelasan section
                if penjelasan_section_pat.match(stripped):
                    # flush any pending markers before leaving main content
                    flush_pasal()
                    in_penjelasan_section = True
                    _logger.info("\n✓ Reached Penjelasan section, stopping insertion")
                    break
                
                # Check for BAB/BAGIAN headers - these should trigger flush of pending pasal
                if bab_header_pat.match(stripped):
                    flush_pasal()
                    result_lines.append(line)
                    continue
                
                # Track current Pasal in main content
                pasal_match = pasal_main_pat.match(stripped)
                if pasal_match:
                    # entering a new pasal: flush previous pending markers
                    if current_pasal is not None:
                        flush_pasal()
                    
                    current_pasal = pasal_match.group(1)
                    current_ayat = None
                    result_lines.append(line)
                    
                    # Prepare Pasal-level penjelasan to be inserted AFTER pasal content
                    key = f'pasal_{current_pasal}'
                    if key in penjelasan_map:
                        pending_pasal_lines = [
                            '',
                            f'💡 [Penjelasan Pasal {current_pasal}]:',
                            f'{penjelasan_map[key]}',
                            ''
                        ]
                        _logger.info(f"  ↪ Queued penjelasan for Pasal {current_pasal}")
                    else:
                        pending_pasal_lines = None
                    continue
                
                # Track current Ayat
                ayat_match = ayat_main_pat.match(stripped)
                if ayat_match and current_pasal:
                    # entering a new ayat: flush previous pending huruf/ayat
                    if current_ayat is not None:
                        flush_ayat()
                    
                    current_ayat = ayat_match.group(1)
                    result_lines.append(line)
                    
                    # Prepare Ayat-level penjelasan to be inserted AFTER ayat content
                    key = f'pasal_{current_pasal}_ayat_{current_ayat}'
                    if key in penjelasan_map:
                        pending_ayat_lines = [
                            '',
                            f'💡 [Penjelasan Ayat ({current_ayat})]:',
                            f'{penjelasan_map[key]}',
                            ''
                        ]
                        _logger.info(f"  ↪ Queued penjelasan for Pasal {current_pasal} Ayat ({current_ayat})")
                    else:
                        pending_ayat_lines = None
                    continue
                
                # Track Huruf (only if inside numbered list/ayat context)
                if current_ayat:
                    huruf_match = huruf_main_pat.match(stripped)
                    if huruf_match:
                        # entering a new huruf: flush previous pending huruf first
                        flush_huruf()
                        huruf = huruf_match.group(1)
                        result_lines.append(line)
                        
                        # Prepare Huruf-level penjelasan to be inserted AFTER huruf content
                        key = f'pasal_{current_pasal}_ayat_{current_ayat}_huruf_{huruf}'
                        if key in penjelasan_map:
                            pending_huruf_lines = [
                                f'   💡 [Penjelasan Huruf {huruf}]:',
                                f'   {penjelasan_map[key]}',
                                ''
                            ]
                            _logger.info(f"  ↪ Queued penjelasan for Pasal {current_pasal} Ayat ({current_ayat}) Huruf {huruf}")
                        else:
                            pending_huruf_lines = None
                        continue
                
                # Regular line - just add it
                result_lines.append(line)
            
            # Flush any remaining pending markers at end of main content
            flush_pasal()
            
            _logger.info(f"\n✓ Penjelasan insertion complete")
            _logger.info("=" * 80)
            
            return '\n'.join(result_lines)
            
        except Exception as e:
            _logger.warning(f"Failed to insert penjelasan: {e}")
            return text

    def _extract_text_from_txt(self, txt_data):
        """Extract and format text from TXT upload (base64-encoded)."""
        try:
            # Decode base64 TXT data to bytes
            txt_bytes = base64.b64decode(txt_data)
            text = None
            # Try common encodings
            for enc in ('utf-8', 'utf-16', 'cp1252', 'latin-1'):
                try:
                    text = txt_bytes.decode(enc)
                    break
                except Exception:
                    continue
            if text is None:
                # Fallback with replacement
                text = txt_bytes.decode('utf-8', errors='replace')

            # Normalize newlines
            text = text.replace('\r\n', '\n').replace('\r', '\n')

            # Insert penjelasan into text if Penjelasan section exists
            # Check for both "PENJELASAN ATAS" and "Penjelasan" (case-insensitive)
            if 'PENJELASAN' in text.upper():
                try:
                    _logger.info("Detected Penjelasan section, starting insertion...")
                    text = self._insert_penjelasan_into_text(text)
                    _logger.info("Successfully inserted penjelasan into text")
                except Exception as e:
                    _logger.error(f"Failed to insert penjelasan: {e}", exc_info=True)
                    _logger.warning(f"Using original text without penjelasan integration")
            else:
                _logger.info("No Penjelasan section found in text")
            
            # Format to HTML using existing formatter
            formatted = self._format_text_to_html(text)
            return f'<div class="txt-content" style="line-height: 1.8;">{formatted}</div>'
        except Exception as e:
            _logger.error(f"Error extracting TXT text: {str(e)}", exc_info=True)
            return f'<p class="text-danger"><strong>Error:</strong> Gagal mengekstrak teks dari TXT: {str(e)}</p>'

    @api.onchange('file_docx')
    def _onchange_file_docx(self):
        """Auto-extract text from DOCX and convert to PDF"""
        if self.file_docx:
            _logger.info(f"DOCX uploaded for regulation {self.id or 'new'}")
            
            # 1. Extract text from DOCX (more accurate)
            extracted_text = self._extract_text_from_docx(self.file_docx)
            if extracted_text and '<strong>Error:</strong>' not in extracted_text:
                self.isi_peraturan = extracted_text
                _logger.info(f"Successfully extracted text from DOCX ({len(extracted_text)} characters)")
            
            # 2. Convert DOCX to PDF for download
            pdf_data = self._convert_docx_to_pdf(self.file_docx)
            if pdf_data:
                self.file_pdf = pdf_data
                _logger.info("Successfully converted DOCX to PDF for download")
        else:
            # File DOCX dihapus - clear content dan PDF
            _logger.info(f"DOCX removed for regulation {self.id or 'new'}, clearing content")
            self.isi_peraturan = '<p>Isi peraturan belum tersedia</p>'
            # Only clear PDF if it was auto-generated from DOCX
            if not self.file_pdf or self.file_pdf:
                # Check if we should clear PDF (don't clear if manually uploaded)
                pass  # Let user manage PDF separately

    @api.onchange('file_txt')
    def _onchange_file_txt(self):
        """Auto-extract text from TXT"""
        if self.file_txt:
            _logger.info(f"TXT uploaded for regulation {self.id or 'new'}")
            extracted_text = self._extract_text_from_txt(self.file_txt)
            if extracted_text and '<strong>Error:</strong>' not in extracted_text:
                self.isi_peraturan = extracted_text
                _logger.info(f"Successfully extracted text from TXT ({len(extracted_text)} characters)")
        else:
            # TXT removed - clear only if no PDF and no DOCX exists
            if not self.file_pdf and not self.file_docx:
                _logger.info(f"TXT removed and no other files present for regulation {self.id or 'new'}, clearing content")
                self.isi_peraturan = '<p>Isi peraturan belum tersedia</p>'
    
    @api.onchange('file_pdf')
    def _onchange_file_pdf(self):
        """Auto-extract text from PDF when uploaded"""
        if self.file_pdf:
            _logger.info(f"PDF uploaded for regulation {self.id or 'new'}, extracting text...")
            extracted_text = self._extract_text_from_pdf(self.file_pdf)
            
            if extracted_text and '<strong>Error:</strong>' not in extracted_text:
                self.isi_peraturan = extracted_text
                _logger.info(f"Successfully extracted text from PDF ({len(extracted_text)} characters)")
            else:
                _logger.warning("Failed to extract text from PDF or PDF contains no text")
                if not self.isi_peraturan or self.isi_peraturan == '<p>Isi peraturan belum tersedia</p>':
                    self.isi_peraturan = extracted_text
        else:
            # File PDF dihapus - clear content jika tidak ada DOCX
            if not self.file_docx:
                _logger.info(f"PDF removed and no DOCX present for regulation {self.id or 'new'}, clearing content")
                self.isi_peraturan = '<p>Isi peraturan belum tersedia</p>'
    
    @api.onchange('bentuk')
    def _onchange_bentuk(self):
        """Auto-generate bentuk_singkat dari bentuk"""
        if self.bentuk:
            # Mapping umum bentuk singkat
            bentuk_mapping = {
                'Undang-Undang': 'UU',
                'Peraturan Pemerintah': 'PP',
                'Peraturan Presiden': 'Perpres',
                'Peraturan Menteri': 'Permen',
                'Peraturan Daerah': 'Perda',
                'Keputusan Presiden': 'Keppres',
                'Keputusan Menteri': 'Kepmen',
                'Instruksi Presiden': 'Inpres',
                'Surat Edaran': 'SE',
            }
            
            for key, value in bentuk_mapping.items():
                if key.lower() in self.bentuk.lower():
                    self.bentuk_singkat = value
                    break
    
    @api.onchange('tanggal_penetapan')
    def _onchange_tanggal_penetapan(self):
        """Auto set tahun dari tanggal penetapan"""
        if self.tanggal_penetapan:
            self.tahun = self.tanggal_penetapan.year
            # Auto set tanggal berlaku jika belum diisi
            if not self.tanggal_berlaku:
                self.tanggal_berlaku = self.tanggal_penetapan
    
    # Computed Fields
    @api.depends('tipe_dokumen')
    def _compute_hierarchy_order(self):
        """Compute hierarchy order based on Indonesian legal hierarchy"""
        hierarchy_mapping = {
            'uud_1945': 1,           
            'tap_mpr': 2,            
            'undang_undang': 3,      
            'perpu': 4,              
            'peraturan_pemerintah': 5,  
            'peraturan_presiden': 6,    
            'keputusan_presiden': 7,    
            'instruksi_presiden': 8,    
            'peraturan_menteri': 9,     
            'keputusan_menteri': 10,    
            'peraturan_daerah': 11,     
            'peraturan_gubernur': 12,   
        }
        
        for record in self:
            record.hierarchy_order = hierarchy_mapping.get(record.tipe_dokumen, 999)

    @api.model
    def _auto_init(self):
        """Auto-init untuk memastikan kompatibilitas field baru"""
        # Call parent init first
        res = super()._auto_init()
        
        # Check and add missing fields for compatibility
        try:
            # List field baru yang mungkin belum ada
            new_fields = [
                ('isi_peraturan', 'TEXT'),
                ('kata_kunci', 'TEXT'), 
                ('ringkasan', 'TEXT'),
                ('hierarchy_order', 'INTEGER')
            ]
            
            for field_name, field_type in new_fields:
                # Check if field exists
                self._cr.execute("""
                    SELECT column_name 
                    FROM information_schema.columns 
                    WHERE table_name=%s AND column_name=%s
                """, (self._table, field_name))
                
                if not self._cr.fetchone():
                    # Add missing field
                    if field_type == 'TEXT':
                        self._cr.execute(f"ALTER TABLE {self._table} ADD COLUMN {field_name} TEXT")
                    elif field_type == 'INTEGER':
                        self._cr.execute(f"ALTER TABLE {self._table} ADD COLUMN {field_name} INTEGER DEFAULT 999")
                    
                    print(f"[COMPATIBILITY] Added missing field: {field_name}")
            
            # Set default values for existing records with NULL values
            self._cr.execute(f"""
                UPDATE {self._table} 
                SET hierarchy_order = 999 
                WHERE hierarchy_order IS NULL
            """)
            
            self._cr.execute(f"""
                UPDATE {self._table} 
                SET isi_peraturan = '<p>Isi peraturan belum tersedia</p>' 
                WHERE isi_peraturan IS NULL OR isi_peraturan = ''
            """)
            
            self._cr.execute(f"""
                UPDATE {self._table} 
                SET kata_kunci = '' 
                WHERE kata_kunci IS NULL
            """)
            
            self._cr.execute(f"""
                UPDATE {self._table} 
                SET ringkasan = 'Ringkasan belum tersedia' 
                WHERE ringkasan IS NULL OR ringkasan = ''
            """)
            
        except Exception as e:
            # Jika ada error, log tapi jangan stop install
            print(f"[COMPATIBILITY WARNING] {str(e)}")
            
        return res

    @api.model
    def create(self, vals_list):
        """Override create untuk validasi dan auto-extract TXT/DOCX/PDF"""
        # Handle both single dict and list of dicts
        if isinstance(vals_list, dict):
            vals_list = [vals_list]
        
        for vals in vals_list:
            # Auto generate bentuk_singkat jika tidak ada
            if vals.get('bentuk') and not vals.get('bentuk_singkat'):
                bentuk = vals['bentuk']
                if 'Menteri' in bentuk:
                    vals['bentuk_singkat'] = 'Permen'
                elif 'Pemerintah' in bentuk:
                    vals['bentuk_singkat'] = 'PP'
                elif 'Presiden' in bentuk:
                    vals['bentuk_singkat'] = 'Perpres'
            
            # Priority: TXT > DOCX > PDF
            if vals.get('file_txt'):
                try:
                    if not vals.get('isi_peraturan'):
                        extracted_text = self._extract_text_from_txt(vals['file_txt'])
                        if extracted_text and '<strong>Error:</strong>' not in extracted_text:
                            vals['isi_peraturan'] = extracted_text
                            _logger.info("Auto-extracted TXT text during creation")
                except Exception as e:
                    _logger.error(f"Failed to process TXT during creation: {e}")
            elif vals.get('file_docx'):
                try:
                    # Extract text from DOCX
                    if not vals.get('isi_peraturan'):
                        extracted_text = self._extract_text_from_docx(vals['file_docx'])
                        if extracted_text and '<strong>Error:</strong>' not in extracted_text:
                            vals['isi_peraturan'] = extracted_text
                            _logger.info("Auto-extracted DOCX text during creation")
                    
                    # Convert DOCX to PDF for download
                    pdf_data = self._convert_docx_to_pdf(vals['file_docx'])
                    if pdf_data:
                        vals['file_pdf'] = pdf_data
                        _logger.info("Auto-converted DOCX to PDF during creation")
                except Exception as e:
                    _logger.error(f"Failed to process DOCX during creation: {e}")
            
            # Fallback: Extract from PDF if uploaded manually
            elif vals.get('file_pdf') and not vals.get('isi_peraturan'):
                try:
                    extracted_text = self._extract_text_from_pdf(vals['file_pdf'])
                    if extracted_text and '<strong>Error:</strong>' not in extracted_text:
                        vals['isi_peraturan'] = extracted_text
                        _logger.info("Auto-extracted PDF text during creation")
                except Exception as e:
                    _logger.error(f"Failed to auto-extract PDF during creation: {e}")
        
        return super(LegalRegulation, self).create(vals_list)
    
    def write(self, vals):
        """Override write untuk auto-extract TXT/DOCX/PDF saat update"""
        # Check if files are being deleted (set to False)
        file_txt_deleted = 'file_txt' in vals and not vals.get('file_txt')
        file_docx_deleted = 'file_docx' in vals and not vals.get('file_docx')
        file_pdf_deleted = 'file_pdf' in vals and not vals.get('file_pdf')
        
        # Handle file deletion - clear content if all files removed
        if (file_txt_deleted or not self.file_txt) and (file_docx_deleted or not self.file_docx) and file_pdf_deleted:
            if 'isi_peraturan' not in vals:
                vals['isi_peraturan'] = '<p>Isi peraturan belum tersedia</p>'
                _logger.info(f"All files deleted for regulation ID {self.id}, clearing content")
        elif file_txt_deleted and not vals.get('file_pdf') and not self.file_pdf and not self.file_docx:
            # TXT deleted and no other file exists
            if 'isi_peraturan' not in vals:
                vals['isi_peraturan'] = '<p>Isi peraturan belum tersedia</p>'
                _logger.info(f"TXT deleted and no other files present for regulation ID {self.id}, clearing content")
        elif file_docx_deleted and not vals.get('file_pdf') and not self.file_pdf and not self.file_txt:
            # DOCX deleted and no PDF/TXT exists
            if 'isi_peraturan' not in vals:
                vals['isi_peraturan'] = '<p>Isi peraturan belum tersedia</p>'
                _logger.info(f"DOCX deleted and no PDF/TXT present for regulation ID {self.id}, clearing content")
        elif file_pdf_deleted and not vals.get('file_docx') and not self.file_docx and not self.file_txt:
            # PDF deleted and no DOCX/TXT exists
            if 'isi_peraturan' not in vals:
                vals['isi_peraturan'] = '<p>Isi peraturan belum tersedia</p>'
                _logger.info(f"PDF deleted and no DOCX/TXT present for regulation ID {self.id}, clearing content")
        
        # Priority: TXT > DOCX > PDF
        if vals.get('file_txt') and not file_txt_deleted:
            try:
                if 'isi_peraturan' not in vals:
                    extracted_text = self._extract_text_from_txt(vals['file_txt'])
                    if extracted_text and '<strong>Error:</strong>' not in extracted_text:
                        vals['isi_peraturan'] = extracted_text
                        _logger.info(f"Auto-extracted TXT text for regulation ID {self.id}")
            except Exception as e:
                _logger.error(f"Failed to process TXT during write: {e}")
        elif vals.get('file_docx') and not file_docx_deleted:
            try:
                # Extract text from DOCX
                if 'isi_peraturan' not in vals:
                    extracted_text = self._extract_text_from_docx(vals['file_docx'])
                    if extracted_text and '<strong>Error:</strong>' not in extracted_text:
                        vals['isi_peraturan'] = extracted_text
                        _logger.info(f"Auto-extracted DOCX text for regulation ID {self.id}")
                
                # Convert DOCX to PDF for download
                pdf_data = self._convert_docx_to_pdf(vals['file_docx'])
                if pdf_data:
                    vals['file_pdf'] = pdf_data
                    _logger.info(f"Auto-converted DOCX to PDF for regulation ID {self.id}")
            except Exception as e:
                _logger.error(f"Failed to process DOCX during write: {e}")
        
        # Fallback: Extract from PDF if uploaded manually
        elif vals.get('file_pdf') and not file_pdf_deleted and 'isi_peraturan' not in vals:
            try:
                extracted_text = self._extract_text_from_pdf(vals['file_pdf'])
                if extracted_text and '<strong>Error:</strong>' not in extracted_text:
                    vals['isi_peraturan'] = extracted_text
                    _logger.info(f"Auto-extracted PDF text for regulation ID {self.id}")
            except Exception as e:
                _logger.error(f"Failed to auto-extract PDF during write: {e}")
        
        return super(LegalRegulation, self).write(vals)
    
    def name_get(self):
        """Custom name display"""
        result = []
        for record in self:
            if record.bentuk_singkat and record.nomor and record.tahun:
                name = f"{record.bentuk_singkat} No. {record.nomor}/{record.tahun}"
            else:
                name = record.judul[:50] + '...' if len(record.judul) > 50 else record.judul
            result.append((record.id, name))
        return result
    
    @api.model
    def _name_search(self, name, args=None, operator='ilike', limit=100, name_get_uid=None):
        """Enhanced search functionality"""
        args = args or []
        domain = []
        if name:
            domain = ['|', '|', '|', '|', '|', '|', '|', '|',
                     ('judul', operator, name),
                     ('nomor', operator, name),
                     ('bentuk_singkat', operator, name),
                     ('subjek', operator, name),
                     ('keterangan', operator, name),
                     ('isi_peraturan', operator, name),
                     ('kata_kunci', operator, name),
                     ('ringkasan', operator, name),
                     ('bentuk', operator, name)]
        
        regulation_ids = self._search(domain + args, limit=limit, access_rights_uid=name_get_uid)
        return self.browse(regulation_ids).name_get()
    
    def action_download_pdf(self):
        """Action untuk download file PDF"""
        self.ensure_one()
        if not self.file_pdf:
            from odoo.exceptions import UserError
            raise UserError('File PDF tidak tersedia untuk peraturan ini.')
        
        return {
            'type': 'ir.actions.act_url',
            'url': f'/web/content?model=legal.regulation&id={self.id}&field=file_pdf&filename_field=file_name&download=true',
            'target': 'self',
        }
    
    def action_reextract_pdf(self):
        """Re-extract content from uploaded file (TXT/DOCX/PDF) with latest extraction method"""
        self.ensure_one()
        
        if not (self.file_txt or self.file_docx or self.file_pdf):
            from odoo.exceptions import UserError
            raise UserError('Tidak ada file untuk di-ekstrak ulang.')
        
        try:
            # Determine which file to extract from (priority: TXT > DOCX > PDF)
            if self.file_txt:
                _logger.info(f"Re-extracting TXT for regulation ID {self.id}: {self.judul}")
                extracted_html = self._extract_text_from_txt(self.file_txt)
                file_type = "TXT"
            elif self.file_docx:
                _logger.info(f"Re-extracting DOCX for regulation ID {self.id}: {self.judul}")
                extracted_html = self._extract_text_from_docx(self.file_docx)
                file_type = "DOCX"
            else:
                _logger.info(f"Re-extracting PDF for regulation ID {self.id}: {self.judul}")
                extracted_html = self._extract_text_from_pdf(self.file_pdf)
                file_type = "PDF"
            
            # Embed extraction timestamp as HTML comment for cache-busting/debugging
            from odoo.tools import format_datetime
            ts = format_datetime(self.env, fields.Datetime.now())
            extracted_html = f"<!-- reextracted from {file_type}: {ts} -->\n" + (extracted_html or '')
            
            # Update the isi_peraturan field
            self.write({
                'isi_peraturan': extracted_html
            })
            
            _logger.info(f"Successfully re-extracted {file_type} for regulation ID {self.id}")
            
            # Reload the view to ensure the latest HTML is shown
            return {'type': 'ir.actions.client', 'tag': 'reload'}
            
        except Exception as e:
            _logger.error(f"Failed to re-extract file for regulation ID {self.id}: {e}")
            from odoo.exceptions import UserError
            raise UserError(f'Gagal mengekstrak ulang file: {str(e)}')

    @api.model
    def _auto_init(self):
        """Ensure compatibility during install - add missing fields safely"""
        res = super()._auto_init()
        
        try:
            # Check dan tambahkan field yang mungkin belum ada
            fields_to_check = [
                ('isi_peraturan', 'TEXT', '<p>Isi peraturan belum tersedia</p>'),
                ('kata_kunci', 'VARCHAR', ''),
                ('ringkasan', 'TEXT', 'Ringkasan belum tersedia'),
                ('hierarchy_order', 'INTEGER', '999')
            ]
            
            for field_name, field_type, default_value in fields_to_check:
                self._cr.execute("""
                    SELECT column_name 
                    FROM information_schema.columns 
                    WHERE table_name='legal_regulation' AND column_name=%s
                """, (field_name,))
                
                if not self._cr.fetchone():
                    if field_type == 'TEXT':
                        self._cr.execute(f"ALTER TABLE legal_regulation ADD COLUMN {field_name} TEXT")
                        if default_value:
                            self._cr.execute(f"UPDATE legal_regulation SET {field_name} = %s WHERE {field_name} IS NULL", (default_value,))
                    elif field_type == 'VARCHAR':
                        self._cr.execute(f"ALTER TABLE legal_regulation ADD COLUMN {field_name} VARCHAR")
                        if default_value:
                            self._cr.execute(f"UPDATE legal_regulation SET {field_name} = %s WHERE {field_name} IS NULL", (default_value,))
                    elif field_type == 'INTEGER':
                        self._cr.execute(f"ALTER TABLE legal_regulation ADD COLUMN {field_name} INTEGER DEFAULT {default_value}")
                        
            self._cr.commit()
            
        except Exception as e:
            # Jika ada error, log tapi jangan crash
            import logging
            _logger = logging.getLogger(__name__)
            _logger.warning(f"Field migration error: {e}")
            
        return res


class LegalRegulationType(models.Model):
    _name = 'legal.regulation.type'
    _description = 'Legal Regulation Type'
    _order = 'sequence, name'
    
    name = fields.Char('Nama Tipe', required=True)
    code = fields.Char('Kode', required=True)
    description = fields.Text('Deskripsi')
    sequence = fields.Integer('Urutan', default=10)
    active = fields.Boolean('Active', default=True)
    
    _sql_constraints = [
        ('code_unique', 'unique(code)', 'Kode tipe peraturan harus unik!')
    ]