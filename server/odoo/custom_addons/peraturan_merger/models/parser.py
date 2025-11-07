# peraturan_merger/models/parser.py
import re
import tempfile
import base64
from io import BytesIO
from difflib import SequenceMatcher
from docx import Document
from unidecode import unidecode

# Tika optional; use if available
try:
    from tika import parser as tika_parser
    TIKA_AVAILABLE = True
except Exception:
    TIKA_AVAILABLE = False

# pdfminer fallback
try:
    from pdfminer.high_level import extract_text
    PDFMINER_AVAILABLE = True
except Exception:
    PDFMINER_AVAILABLE = False

class DocParser:
    """
    Parser untuk mengekstrak struktur pasal/ayat dari PDF/DOCX.
    Fungsi utama: parse_bytes(filebytes, filename) -> list of articles
    """

    # dasar regex; dapat disesuaikan untuk format instansi
    ARTICLE_RE = re.compile(r'(?i)^\s*Pasal\s+(\d+)\b', re.MULTILINE)
    AYAT_RE = re.compile(r'(?i)^\s*ayat\s*\(?\s*(\d+[a-zA-Z]?)\s*\)?', re.MULTILINE)
    # juga cari numbering seperti "(1)" di awal baris
    AYAT_BRACKET_RE = re.compile(r'(?m)^\s*\(?\s*(\d+[a-zA-Z]?)\s*\)\s*')

    def parse_bytes(self, filebytes, filename):
        """
        Menerima raw bytes (b'...') dari file, kembalikan struktur articles.
        """
        text = None
        # try tika
        if TIKA_AVAILABLE:
            try:
                tmp = tempfile.NamedTemporaryFile(delete=False)
                tmp.write(filebytes)
                tmp.close()
                parsed = tika_parser.from_file(tmp.name)
                text = parsed.get('content') or ''
            except Exception:
                text = None
        if not text and PDFMINER_AVAILABLE and filename.lower().endswith('.pdf'):
            try:
                # pdfminer expects filename or bytes stream; write tmp file
                tmp = tempfile.NamedTemporaryFile(delete=False)
                tmp.write(filebytes)
                tmp.close()
                text = extract_text(tmp.name)
            except Exception:
                text = None
        if not text and filename.lower().endswith('.docx'):
            try:
                # try python-docx
                doc = Document(BytesIO(filebytes))
                parts = []
                for p in doc.paragraphs:
                    parts.append(p.text)
                text = '\n'.join(parts)
            except Exception:
                text = None
        if not text:
            # fallback - try decode as utf-8
            try:
                text = filebytes.decode('utf-8', errors='ignore')
            except Exception:
                text = ''

        text = unidecode(text)
        return self._parse_text_to_structure(text)

    def parse_text_fallback(self, text):
        text = unidecode(text)
        return self._parse_text_to_structure(text)

    def _parse_text_to_structure(self, text):
        articles = []
        # Split by occurrences of "Pasal X"
        # We keep the 'Pasal X' line with each chunk
        parts = re.split(r'(?i)(?=^\s*Pasal\s+\d+\b)', text, flags=re.MULTILINE)
        for part in parts:
            m = self.ARTICLE_RE.search(part)
            if not m:
                continue
            no = m.group(1).strip()
            # article title: first non-empty line after Pasal X (if any)
            lines = [ln.strip() for ln in part.splitlines() if ln.strip()]
            article_title = ''
            if lines:
                # lines[0] likely contains "Pasal X ..." — find remainder after that
                first = lines[0]
                # if "Pasal X" and title in same line
                mm = re.match(r'(?i)Pasal\s+' + re.escape(no) + r'\b(.*)', first)
                if mm:
                    article_title = mm.group(1).strip()
                    if not article_title and len(lines) > 1:
                        article_title = lines[1].strip()
                else:
                    article_title = lines[0].strip()

            # find ayat blocks inside part
            paragraphs = []
            # split by 'ayat' or bracket numbering at line starts
            # Strategy: find all positions of ayat headings and slice
            ayat_positions = []
            for m_ayat in self.AYAT_RE.finditer(part):
                ayat_positions.append((m_ayat.start(), m_ayat.group(1)))
            # fallback to bracket style "(1)"
            if not ayat_positions:
                for m_ab in self.AYAT_BRACKET_RE.finditer(part):
                    ayat_positions.append((m_ab.start(), m_ab.group(1)))
            # if we found headings, slice segment-wise
            if ayat_positions:
                ayat_positions.sort()
                for idx, (pos, label) in enumerate(ayat_positions):
                    start = pos
                    end = ayat_positions[idx+1][0] if idx+1 < len(ayat_positions) else len(part)
                    snippet = part[start:end].strip()
                    # remove heading token
                    snippet = re.sub(r'(?i)^\s*(ayat\s*\(?\s*' + re.escape(label) + r'\s*\)?|\(?\s*' + re.escape(label) + r'\s*\)?)\s*', '', snippet).strip()
                    paragraphs.append({'label': label, 'text': snippet, 'order': idx+1})
            else:
                # no ayat headings: take remainder after pasal heading as one paragraph
                body = re.sub(r'(?i)^\s*Pasal\s+' + re.escape(no) + r'\b.*', '', part, flags=re.DOTALL).strip()
                if body:
                    paragraphs.append({'label': '', 'text': body, 'order': 1})
            articles.append({'article_no': no, 'article_title': article_title, 'paragraphs': paragraphs})
        return articles

def _diff_text(a, b):
    return SequenceMatcher(None, a or '', b or '').ratio()

def merge_versions_to_consolidation(record, versions, mode='annotated'):
    """
    Kumpulkan struktur per artikel; buat dokumen DOCX yang berisi consolidated body.
    mode: 'annotated' (default) | 'final' | 'history'
    Mengembalikan bytes DOCX.
    """
    # Build mapping: article_no -> {version_index: [ (label, text, seq) ] }
    articles_map = {}
    for idx, v in enumerate(versions):
        for art in v.article_ids:
            key = art.number or f"unknown_{art.id}"
            if key not in articles_map:
                articles_map[key] = {}
            paras = []
            # ensure paragraph ordering
            for p in art.paragraph_ids.sorted(key=lambda r: r.sequence):
                paras.append((p.label or '', p.text or '', p.sequence))
            articles_map[key][idx] = paras

    # Prepare DOCX
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.add_heading(f'Badan Peraturan Terpadu: {record.rule_id.name}', level=1)
    versions_line = ', '.join([f"{v.name} ({v.tanggal or v.create_date.date()})" for v in versions])
    doc.add_paragraph(f'Versi digabung: {versions_line}')

    # iterate articles sorted numerically if possible
    def art_key_sort(k):
        try:
            return int(re.sub(r'\D','',k) or 0)
        except Exception:
            return 999999
    for art_key in sorted(articles_map.keys(), key=art_key_sort):
        doc.add_heading(f'Pasal {art_key}', level=2)
        # collect all paragraph labels across versions, preserving order by first appearance
        label_order = []
        for vidx in sorted(articles_map[art_key].keys()):
            for label, text, seq in articles_map[art_key][vidx]:
                if label not in label_order:
                    label_order.append(label)

        # for each label build list of texts across versions
        for label in label_order:
            texts_by_version = []
            for vidx in range(len(versions)):
                paras = articles_map[art_key].get(vidx, [])
                t = None
                for lab, txt, seq in paras:
                    if lab == label:
                        t = txt
                        break
                texts_by_version.append((vidx, t))
            final_text = texts_by_version[-1][1]
            p = doc.add_paragraph()
            if label:
                run_lbl = p.add_run(f'({label}) ')
                run_lbl.bold = True
            # populate based on mode
            if mode == 'final':
                p.add_run(final_text or '[-- Dihapus --]')
            elif mode == 'history':
                for vidx, t in texts_by_version:
                    v = versions[vidx]
                    p.add_run(f'\nVersi {v.name} ({v.tanggal or v.create_date.date()}): ').bold = True
                    p.add_run(t or '[kosong/dihapus]')
            else:  # annotated
                if final_text:
                    p.add_run(final_text)
                else:
                    p.add_run('[Dihapus pada versi terakhir]').italic = True
                # add change notes
                notes = []
                for vidx, t in texts_by_version[:-1]:  # compare earlier versions to last
                    if t is None and final_text:
                        notes.append(f'Ditambahkan (sebelum: kosong) — sumber: {versions[-1].name}')
                    elif t and not final_text:
                        notes.append(f'Dihapus (sebelumnya ada pada {versions[vidx].name})')
                    elif t and final_text and _diff_text(t, final_text) < 0.95:
                        notes.append(f'Perubahan antara {versions[vidx].name} → {versions[-1].name}')
                if notes:
                    p.add_run('\nCatatan perubahan: ' + '; '.join(notes)).italic = True

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()
